import docx
import json
import re

def extract_all_tables_data_from_docx(filename):
    """
    从 docx 文件中提取所有表格的数据
    """
    doc = docx.Document(filename)
    all_data = []
    for table_index, table in enumerate(doc.tables):
        print(f"正在处理第 {table_index + 1} 个表格...")
        rows_in_this_table = 0
        for i, row in enumerate(table.rows):
            if i == 0: # 跳过每个表格的表头
                continue
            cells = [cell.text.strip() for cell in row.cells]
            # 假设列顺序为：序号、单词、音标、释义
            if len(cells) >= 4:
                raw_num = cells[0]
                raw_word = cells[1]
                raw_symbol = cells[2] # 音标已包含 /
                raw_trans = cells[3]
                all_data.append((raw_num, raw_word, raw_symbol, raw_trans))
                rows_in_this_table += 1
            else:
                print(f"  警告：第 {table_index + 1} 个表格第 {i+1} 行数据不足4列，已跳过: {cells}")
        print(f"  第 {table_index + 1} 个表格处理完成，提取了 {rows_in_this_table} 行数据。")
    print(f"总共从 {len(doc.tables)} 个表格中提取到 {len(all_data)} 行数据。")
    return all_data

def process_cet6_data(data):
    """
    处理提取的数据，生成 JSON 对象列表
    """
    json_objects = []
    for row_data in data:
        raw_num, raw_word, raw_symbol, raw_trans = row_data

        # 处理 num
        num_value = str(raw_num).strip()
        # 移除可能的非数字字符，只保留数字部分 (例如 '4618' from '4618' 或 '4618*')
        num_match = re.search(r'\d+', num_value)
        if num_match:
            clean_num = num_match.group(0)
        else:
            print(f"警告：无法从 '{num_value}' 中提取数字，跳过此行。")
            continue
        num = f"CET6-{clean_num}"

        # 处理 word
        word = str(raw_word).strip()

        # 处理 symbol (音标已包含 /，直接使用)
        symbol = str(raw_symbol).strip()
        # 可选：清理音标中的多余空格，例如 '/ ə ˈ bænd ə n/' -> '/əˈbændən/'
        # symbol = re.sub(r'\s+', '', symbol)
        # 或者更精确地只处理 / 和音素之间的空格
        symbol_cleaned = re.sub(r'/\s+|(?<=\S)\s+/', '', symbol) # 移除 / 前后和音素间的空格
        # print(f"Original: '{raw_symbol}', Cleaned: '{symbol_cleaned}'") # 调试用，可删除
        symbol = symbol_cleaned

        # 处理 trans
        trans = str(raw_trans).strip()

        item = {
            "num": num,
            "word": word,
            "symbol": symbol,
            "trans": trans
        }
        json_objects.append(item)
    return json_objects

# --- 主执行流程 ---
if __name__ == "__main__":
    input_filename = '大学英语六级词汇表(全)含音标.docx'

    try:
        # 1. 从 docx 提取所有表格数据
        print(f"正在从 {input_filename} 提取所有表格数据...")
        raw_data = extract_all_tables_data_from_docx(input_filename)

        # 2. 处理数据并生成 JSON 对象
        print("正在处理数据并生成 JSON...")
        json_list = process_cet6_data(raw_data)

        # 3. 输出 JSON
        if json_list:
            # 输出为 JSON 格式字符串（每项一行，便于复制粘贴）
            json_lines = [json.dumps(item, ensure_ascii=False) for item in json_list]
            output = ",\n".join(json_lines)

            # 打印结果
            print("[")
            print(output)
            print("]")

            # 可选：保存到文件
            output_filename = 'cet6_output_full.json'
            with open(output_filename, 'w', encoding='utf-8') as f:
                f.write("[\n")
                f.write(",\n".join(json_lines))
                f.write("\n]")
            print(f"\n结果已保存到 {output_filename}")
            print(f"总共生成了 {len(json_list)} 个词汇的 JSON 对象。")
        else:
            print("处理后没有生成任何有效的 JSON 对象。")

    except FileNotFoundError:
        print(f"错误：找不到文件 {input_filename}")
    except Exception as e:
        print(f"处理过程中发生错误: {e}")
